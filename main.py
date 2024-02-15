import tkinter as tk 
from tkinter import filedialog
import customtkinter as ctk
from CTkMessagebox import CTkMessagebox
import re 
import docx 
import openpyxl 
from PyPDF2 import PdfReader 
import csv 
import requests 
from bs4 import BeautifulSoup 
import os 
import webbrowser 
import pyperclip
import threading

def createToolTip(widget, text):
    try:
        def enter(event):
            widget._after_id = widget.after(600, show_tooltip, event)

        def leave(event):
            widget.after_cancel(widget._after_id)
            tooltip = getattr(widget, "_tooltip", None)
            if tooltip:
                tooltip.destroy()
                widget._tooltip = None

        def show_tooltip(event):
            tooltip = ctk.CTkToplevel(widget)
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            label = ctk.CTkLabel(tooltip, text=text, fg_color="black", text_color="white")
            label.grid()
            widget._tooltip = tooltip

        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)
    except Exception as e:
        CTkMessagebox(title="Error", message=f"An error occurred: {e}")

class App(ctk.CTk):
    def __init__(self):
        try:
            super().__init__()

            self.title('Word Counter Master')
            self.iconbitmap('Word Counter Master.ico')
            self.resizable(False, False)

            self.main_frame = ctk.CTkFrame(self)
            self.main_frame.grid(row=0, column=0, sticky="nsew")

            self.left_frame = ctk.CTkFrame(self.main_frame, width=200, height=400)
            self.left_frame.grid(row=0, column=0, padx=10, pady=5, sticky="ns")

            self.right_frame = ctk.CTkFrame(self.main_frame)
            self.right_frame.grid(row=0, column=1, padx=10, pady=5, sticky="nsew")

            self.text_input = ctk.CTkTextbox(self.right_frame, font=("Arial", 12), width=720, height=430)
            self.text_input.grid(row=0, column=0, sticky="nsew")
            createToolTip(self.text_input, "Enter the text you want to count here")
            self.text_input.bind("<KeyRelease>", self.count_words)

            self.label_count = ctk.CTkLabel(self.right_frame, text="", font=("Arial", 10))
            self.label_count.grid(row=1, column=0, sticky="ns")
            createToolTip(self.label_count, "This displays the word count")
            self.count_words()

            self.frame = ctk.CTkFrame(self.right_frame)
            self.frame.grid(row=2, column=0, sticky="ew", pady=10)
            self.search_entry = ctk.CTkEntry(self.frame, width=600, height=30)
            self.search_button = ctk.CTkButton(self.frame, width=15, text="Search", command=self.search_text)
            self.frame.grid_columnconfigure(0, weight=1)
            self.search_entry.grid(row=0, column=0, sticky="w", padx=5, ipadx=10)
            self.search_button.grid(row=0, column=1, sticky="e", padx=5, ipadx=10)
            createToolTip(self.search_entry, "Enter text to search")
            createToolTip(self.search_button, "Click to search")

            self.open_button = ctk.CTkButton(self.left_frame, text="Open File", command=lambda: threading.Thread(target=self.open_file).start())
            self.open_button.grid(row=0, column=0, sticky="ew", pady=10)
            createToolTip(self.open_button, "Click to open a file")

            self.save_button = ctk.CTkButton(self.left_frame, text="Save Text", command=self.save_file)
            self.save_button.grid(row=1, column=0, sticky="ew", pady=10)
            createToolTip(self.save_button, "Click to save the text")

            self.copy_button = ctk.CTkButton(self.left_frame, text="Copy Text", command=self.copy_text)
            self.copy_button.grid(row=2, column=0, sticky="ew", pady=10)
            createToolTip(self.copy_button, "Click to copy the text")

            self.clear_button = ctk.CTkButton(self.left_frame, text="Clear Text", command=self.clear_text)
            self.clear_button.grid(row=3, column=0, sticky="ew", pady=10)
            createToolTip(self.clear_button, "Clear the text box")

            self.fetch_button = ctk.CTkButton(self.left_frame, text="Fetch Web Content", command=self.fetch_web_content)
            self.fetch_button.grid(row=4, column=0, sticky="ew", pady=10)
            createToolTip(self.fetch_button, "Click to fetch web content")

            self.clear_copy_button = ctk.CTkButton(self.left_frame, text="Clear and Copy to Clipboard", command=self.clear_and_copy_to_clipboard)
            self.clear_copy_button.grid(row=5, column=0, sticky="ew", pady=10)
            createToolTip(self.clear_copy_button, "Clear the text box and copy its contents to clipboard")

            self.feedback_button = ctk.CTkButton(self.left_frame, text="Send Feedback", command=self.open_feedback_link)
            self.feedback_button.grid(row=7, column=0, sticky="ew", pady=10)
            createToolTip(self.feedback_button, "Click to send feedback")

            self.shortcuts_button = ctk.CTkButton(self.left_frame, text="View Shortcuts", command=self.view_shortcuts)
            self.shortcuts_button.grid(row=6, column=0, sticky="ew", pady=10)
            createToolTip(self.shortcuts_button, "Click to view all shortcuts")

            self.main_frame.columnconfigure(1, weight=1)
            self.main_frame.rowconfigure(0, weight=1)

            self.right_frame.columnconfigure(0, weight=1)
            self.right_frame.rowconfigure(0, weight=1)

            self.bind('<Control-o>', lambda event: self.open_file())
            self.bind('<Control-s>', lambda event: self.save_file())
            self.bind('<Control-c>', lambda event: self.copy_text())
            self.bind('<Control-x>', lambda event: self.clear_text())
            self.bind('<Control-f>', lambda event: self.fetch_web_content())
            self.bind('<Control-e>', lambda event: self.open_feedback_link())
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def open_file(self, loading_window=None):
        try:
            file_path = filedialog.askopenfilename(filetypes=[('Text Documents', '*.txt *.docx *.doc *.xlsx *.xls *.pdf *.csv')])
            if file_path:
                if file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()

                elif file_path.endswith('.doc') or file_path.endswith('.docx'):
                    doc = docx.Document(file_path)
                    text = '\n'.join([para.text for para in doc.paragraphs])

                elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
                    wb = openpyxl.load_workbook(file_path)
                    text = ""
                    for sheet in wb:
                        for row in sheet.rows:
                            for cell in row:
                                if isinstance(cell.value, str):
                                    text += cell.value
                                else:
                                    text += str(cell.value)

                elif file_path.endswith('.pdf'):
                    pdfFileObj = open(file_path, 'rb')
                    pdfReader = PdfReader(pdfFileObj)
                    text = ""
                    for pageNum in range(len(pdfReader.pages)):
                        pageObj = pdfReader.pages[pageNum]
                        text += pageObj.extract_text()
                    pdfFileObj.close()

                elif file_path.endswith('.csv'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        reader = csv.reader(f)
                        text = '\n'.join([' '.join(row) for row in reader])

                self.text_input.delete(1.0, "end")
                self.text_input.insert(1.0, text)
                self.count_words()
                if loading_window:
                    loading_window.destroy()
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def fetch_web_content(self):
        url_dialog = ctk.CTkInputDialog(title="Input", text="Enter the URL")
        url = url_dialog.get_input()
        if url:
            try:
                response = requests.get(url)
                soup = BeautifulSoup(response.text, 'html.parser')
                text = soup.get_text()
                self.text_input.delete('1.0', tk.END)
                self.text_input.insert(tk.END, text)
                self.count_words()
            except Exception as e:
                CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def search_text(self):
        try:
            search_text = self.search_entry.get()
            if search_text:
                start_pos = "1.0"
                end_pos = tk.END
                self.text_input.tag_remove("search", "1.0", tk.END)
                while True:
                    start_pos = self.text_input.search(search_text, start_pos, stopindex=end_pos, nocase=True)
                    if not start_pos:
                        break
                    end_pos = f"{start_pos}+{len(search_text)}c"
                    self.text_input.tag_add("search", start_pos, end_pos)
                    self.text_input.see(start_pos)
                    start_pos = end_pos
                self.text_input.tag_config("search", background="yellow", foreground="black")
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def count_words(self, event=None):
        try:
            text = self.text_input.get("1.0", "end")
            word_counts = self.count_all_words(text)
            self.label_count.configure(text=word_counts)
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def count_all_words(self, text):
        try:
            counts = {
                "Chinese Count": len(re.findall(r'[\u4E00-\u9FFF]', text)),
                "English Count": len(re.findall(r'\b[a-zA-Z]+\b', text)),
                "Malay Count": len(re.findall(r'\b[abcçdefgğhıijklmnoöprsştuüvyz]+([\'-][abcçdefgğhıijklmnoöprsştuüvyz]+)*\b', text, re.IGNORECASE)),
                "Number Count": len(re.findall(r'\d+', text)),
                "Comma Count": len(re.findall(r',|，', text)),
                "Period Count": len(re.findall(r'\.|。', text)),
            }
            return "\n".join(f"{key}: {value}" for key, value in counts.items())
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def save_file(self):
        try:
            text = self.text_input.get("1.0", "end").strip()
            if not text:
                CTkMessagebox("Info", "Input Text In The Text Box First!")
                return

            file_name = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[('Text Documents', '*.txt'), ('Word Documents', '*.docx'), ('PDF Files', '*.pdf')], initialfile='textbox')
            if file_name:
                extension = os.path.splitext(file_name)[1].lower()

                if extension == '.txt':
                    with open(file_name, 'w', encoding='utf-8') as f:
                        f.write(text)

                elif extension == '.docx':
                    doc = docx.Document()
                    doc.add_paragraph(text)
                    doc.save(file_name)

                elif extension == '.pdf':
                    from reportlab.platypus import SimpleDocTemplate, Paragraph
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.lib.pagesizes import letter
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont

                    pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))  
                    pdf = SimpleDocTemplate(file_name, pagesize=letter)
                    styles = getSampleStyleSheet()
                    style = styles['Normal']
                    style.fontName = 'SimSun'  
                    style.fontSize = 15
                    story = [Paragraph(text, style)]
                    pdf.build(story)

                else:
                    CTkMessagebox(title="Warning", message="Unsupported file format")
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def open_feedback_link(self):
        webbrowser.open("https://github.com/fatherxtreme123/Word.Counter.Master/issues")

    def copy_text(self):
        try:
            text = self.text_input.get("1.0", "end").strip()
            if text:
                pyperclip.copy(text)
            else:
                CTkMessagebox(title="Info", message="Text box is empty.")
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def clear_and_copy_to_clipboard(self):
        try:
            text = self.text_input.get("1.0", "end").strip()
            if text:
                self.text_input.delete('1.0', tk.END)
                pyperclip.copy(text)
            else:
                CTkMessagebox(title="Info", message="Text box is already empty.")
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def clear_text(self):
        try:
            if self.text_input.get("1.0", "end").strip():
                self.text_input.delete('1.0', tk.END)
                self.count_words()
            else:
                CTkMessagebox(title="Info", message="Text box is already empty.")
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def back_to_main_page(self):
        try:
            self.right_frame.grid()

            if hasattr(self, 'shortcuts_frame'):
                self.shortcuts_frame.destroy()
                
            self.shortcuts_button.configure(text="View Shortcuts", command=self.view_shortcuts)
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

    def view_shortcuts(self):
        try:
            shortcuts_info = (
                "Keyboard Shortcuts:\n"
                "Ctrl + O: Open File\n"
                "Ctrl + S: Save Text\n"
                "Ctrl + C: Copy Text\n"
                "Ctrl + X: Clear Text\n"
                "Ctrl + F: Fetch Web Content\n"
                "Ctrl + E: Send Feedback"
            )
            self.shortcuts_frame = ctk.CTkFrame(self.main_frame)
            self.shortcuts_frame.grid(row=0, column=1, padx=10, pady=5, sticky="nsew")

            self.right_frame.grid_remove()

            self.shortcuts_button.configure(text="Back to Main Page", command=self.back_to_main_page)

            self.shortcuts_label = ctk.CTkLabel(self.shortcuts_frame, text=shortcuts_info)
            self.shortcuts_label.grid(padx=10, pady=10)
        except Exception as e:
            CTkMessagebox(title="Error", message=f"An error occurred: {e}")

if __name__ == "__main__":
    try:
        app = App()
        app.mainloop()
    except Exception as e:
        CTkMessagebox(title="Error", message=f"An error occurred: {e}")
